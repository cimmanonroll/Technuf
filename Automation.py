import streamlit as st
import fitz  # PyMuPDF
import spacy
from langchain_community.llms import Ollama
import tempfile
from docx import Document
import re
import io
 
nlp = spacy.load("en_core_web_sm")
 
llm = Ollama(base_url = "http://10.1.0.133:11434", model="llama3")
 
def get_job_title(text):
    doc = nlp(text)
    job_title = []
    capture = False
    for sent in doc.sents:
        if "title" in sent.text.lower():
            capture = True
        if capture:
            job_title = sent.text.strip()
            break
    job_title = job_title.split("\n")
    index = job_title.index("Task Order Title: ")
    return job_title[index + 1]

# Extracts the text from the .docx file
# Sanitize resume for security (?)
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
 
def extract_text_from_pdf(file_path):
    pdf_document = fitz.open(file_path)
    text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    return text
 
def extract_relevant_sections(text):
    # Extract everything before "Scope of Work" and the required sections
    pattern = re.compile(r'(Scope of Work.*?)(Experience.*?|)(Additionally.*?|)(?=\nDeliverables|$)', re.DOTALL | re.IGNORECASE)
    match = pattern.search(text)
    if match:
        scope_of_work = match.group(1).strip()
        skills_experience = match.group(2).strip()
        additionally_successful = match.group(3).strip()
        return scope_of_work + "\n\n" + skills_experience + "\n\n" + additionally_successful
    return ""

# Generate interview questions specifically for a resume and job description
def generate_interview_questions(job_description, resume):
    prompt = f"""
    You are an interviewer at Technuf LLC that needs to curate interview questions to ask a candidate based on their resume and a given job description.
    According to this job description:
    {job_description}. 
    Generate 5 interivew questions tailored for this resume in bullet point format: 
    {resume}
    """
   
    response = llm.invoke(prompt)
    return response

# Generate job description using only the relevant sections
def generate_job_description(text):
    prompt = f"""
    Generate a job posting for Technuf LLC based in Maryland, USA. Use the information from the following text to create the job description using the template below:
   
    Who we are [Do not change this header]:
    Technuf, LLC is a Maryland based SBA certified 8(a) small business company providing leading-edge and proven technologies, industry vertical domain expertise and highly skilled and motivated professionals to achieve our customers' mission critical business needs.
   
    What we’re looking for [Do not change this header]:
    [Insert a one paragraph description of the name of the job that Technuf is looking for and what the job entails (make sure to start the paragraph with: Technuf is looking for...)]
   
    Responsibilities [Do not change this header]:
    [This section should copy everything from the Scope of Work section]
   
    Skills and Experience [Do not change this header]:
    [This section should copy everything from the Skills/Experience section]
   
    Good to Have [Do not change this header]:
    [This section should copy everything from the Additionally, to be successful section]
   
    Education [Do not change this header]:
    Required: [insert required level of education here]
    Preferred: [insert preferred level of education here]
   
    Technuf is an Equal Opportunity/Affirmative Action Employer. Members of ethnic minorities, women, special disabled veterans, veterans of the Vietnam-era, recently separated veterans, and other protected veterans, persons of disability and/or persons age 40 and over are encouraged to apply.
   
    Ensure to exclude any mention of Evaluation Criteria, How to Apply, Note, and Interview Criteria.
 
    Use the information from the text below to fill in the relevant sections:
   
    {text}
    """
   
    response = llm.invoke(prompt)
    return response

def fill_document(template_path, output_path, data):
    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)
                    
    doc.save(output_path)
    bio = io.BytesIO()
    doc.save(bio)
    if doc:
        st.download_button(
            label="Click here to download",
            data=bio.getvalue(),
            file_name="filled.docx",
            mime="docx"
        )

# TODO: Button for unformatted proposal response, formatted proposal response, input resume to format it
# Focus on input resume + linkedin for now
def main():
    st.title("Technuf Automation")
    st.write("Upload a PDF file")
    data = {}
    solicitation_file = st.file_uploader("Choose a solicitation PDF file", type="pdf")
    if solicitation_file:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(solicitation_file.read())
            temp_filepath = temp_file.name
 
        pdf_text = extract_text_from_pdf(temp_filepath)
        relevant_sections = extract_relevant_sections(pdf_text)

        if relevant_sections:
            st.subheader("Extracted Relevant Sections")
            st.write(relevant_sections)
            if st.button("Generate Job Description"):
                st.session_state.job_button = True
                with st.spinner('Generating job posting'):
                    job_posting = get_job_title(pdf_text) + "\n\n"
                    job_description = generate_job_description(relevant_sections)
                    st.subheader("Generated Job Description")
                    job_posting += job_description
                    job_posting += "\n\nBenefits:\n\nWe offer a competitive pay and benefits package that includes generous paid-time-off including holidays, short-and-long-term disability; group health insurance including medical, dental and vision coverage, training and 401(k) retirement plan. \n"
                    job_posting += "\nTechnuf is an Equal Opportunity/Affirmative Action Employer. Members of ethnic minorities, women, special disabled veterans, veterans of the Vietnam-era, recently separated veterans, and other protected veterans, persons of disability and/or persons age 40 and over are encouraged to apply.\n"
                    newline_index = job_posting.find("Preferred")
                    if newline_index:
                        job_posting = job_posting[:newline_index] + "\n" + job_posting[newline_index:]
                    if 'job_posting' not in st.session_state:
                        st.session_state['job_posting'] = job_posting
                    st.write(job_posting)
                    
                    
        else:
            st.error("No relevant sections found in the uploaded PDF.")
        resume_file = st.file_uploader("Select the resume", type=[".pdf", ".docx"], accept_multiple_files=False)
        if st.button("Generate Interview Questions"):
            if resume_file:
                if resume_file.name.endswith('.docx'):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                        temp_file.write(resume_file.getbuffer())
                        resume_temp_file_path = temp_file.name
                    resume = extract_text_from_docx(resume_temp_file_path)
                elif resume_file.name.endswith('.pdf'):
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                        temp_file.write(resume_file.getbuffer())
                        resume_temp_file_path = temp_file.name
                    resume = extract_text_from_pdf(resume_temp_file_path)
                interview_questions = generate_interview_questions(st.session_state['job_posting'], resume)
                st.write(interview_questions)
                    
        proposal_template = st.file_uploader("Choose a proposal template docx file", type="docx")
        resume_upload_btn = st.button("Input Resumes")
        if "uploadbtn_state" not in st.session_state:
            st.session_state.uploadbtn_state = False
        
        if resume_upload_btn or st.session_state.uploadbtn_state:
            st.session_state.uploadbtn_state = True
            template_file = st.file_uploader("Select the propsal template", type=[".docx"], accept_multiple_files=False)
            resume_files = st.file_uploader("Select the resumes", type=[".pdf", ".docx"], accept_multiple_files=True)
            resumes = []
            if resume_files and template_file:
                st.header("Resume Review")
                for resume in resume_files:
                    if resume.name.endswith('.docx'):
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                            temp_file.write(resume.getbuffer())
                            resume_temp_file_path = temp_file.name
                        resumes.append(extract_text_from_docx(resume_temp_file_path))
                    elif resume.name.endswith('.pdf'):
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                            temp_file.write(resume.getbuffer())
                            resume_temp_file_path = temp_file.name
                        resumes.append(extract_text_from_pdf(resume_temp_file_path))
            
                resume_count = 0
                for resume in resumes:
                    if resume_count == 0:
                        data = {
                            '[Resume1]' : resume
                        }
                    elif resume_count == 1:
                        data['[Resume2]'] = resume
                    resume_count += 1
                fill_document(template_file, 'filled.docx', data)
 
if __name__ == "__main__":
    main()
