import streamlit as st
import fitz  # PyMuPDF
import spacy
import tempfile
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
from docxtpl import DocxTemplate
import streamlit_authenticator as stauth
from streamlit_authenticator.utilities.hasher import Hasher
import re
import os
import io
import json
import yaml
from yaml.loader import SafeLoader
import mammoth
import pdfkit
from streamlit_pdf_viewer import pdf_viewer

load_dotenv()
os.environ["OPENAI_API_KEY"] = os.getenv('OPENAI_API_KEY')
client = OpenAI()
nlp = spacy.load("en_core_web_sm")
def get_job_title_pdf(text):
    # Extracts the job title from a PDF document
    doc = nlp(text)
    job_title = []
    capture = False
    for sent in doc.sents:
        if "title" in sent.text.lower():
            capture = True
        if capture:
            job_title = sent.text.strip()
            break
    job_title = job_title.split("\n")
    for i in range(len(job_title)):
        job_title[i] = job_title[i].strip()
    index = job_title.index("Task Order Title:")
    return job_title[index + 1]

def get_job_title_docx(text):
    # Extracts the job title from a docx document
    doc = nlp(text)
    job_title = []
    capture = False
    for sent in doc.sents:
        if "title" in sent.text.lower():
            capture = True
        if capture:
            job_title = sent.text.strip()
            break
    job_title = job_title.split(": ")
    for i in job_title:
        if i.find("Task Order Title") :
            new_line_index = i.rfind("\n")
            return i[new_line_index + 18:].strip()    
    return "No job title found"

# Sanitize resume for security (?), using a delimiter (############) and removing all instances of "prompt" and "ignore"
def extract_text_from_docx(docx_path):
    # Extracts all text from a docx file into a string
    doc = Document(docx_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

def extract_text_from_pdf(file_path):
    # Extracts all text from a PDF file into a string
    pdf_document = fitz.open(file_path)
    text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    return text

def extract_relevant_sections(text):
    # Extracts "Scope of Work" and other required sections given a text string
    pattern = re.compile(r'(Scope of Work.*?)(Experience.*?|)(Additionally.*?|)(?=\nDeliverables|$)', re.DOTALL)
    match = pattern.search(text)
    if match:
        scope_of_work = match.group(1).strip()
        skills_experience = match.group(2).strip()
        additionally_successful = match.group(3).strip()
        return scope_of_work + "\n\n" + skills_experience + "\n\n" + additionally_successful
    return ""

def generate_resume_descriptions(resume, job_description):
    # Generates a resume description using OpenAI API for a candidate given a resume string and the job description string
    instruction = f""" 
    Generate three paragraphs describing the job candidate's ability to perform their job well given a resume delimited by double quotes and a job description delimited by double quotes. 
    You will only reply in the following format based on the example given, do not include any unecessary greetings or introductions: 
    Mr. Aizaz Ur Rahman is a dynamic and results-oriented Kronos Configuration Consultant with
    over seven (7) years of hands-on experience specializing in Kronos Workforce Central and
    Dimensions. He demonstrates expertise in system administration, user management, and
    configuring modules such as timekeeping, scheduling, and labor analytics. Mr. Rahman is skilled
    in creating and optimizing SQL queries and stored procedures in MSSQL to support complex
    timekeeping processes and reporting requirements. He possesses an extensive background in
    integrating Kronos Time &amp; Attendance with HR and payroll systems using Kronos APIs and
    interfaces, ensuring seamless data flow and accuracy across platforms. Proficient in utilizing Pro
    WFM (Dimensions) and Boomi platforms to enhance system functionalities and support
    operational needs.
    Mr. Rahman has demonstrated success in managing full lifecycle upgrades of Kronos systems,
    from planning and configuration to testing and deployment, ensuring minimal disruption and
    maximum efficiency. He is adept at creating and modifying system configurations, including pay
    rules, work rules, and assignment rules, to align with evolving requirements. Mr. Rahman
    possesses strong analytical skills that he applies to troubleshoot system errors, resolve technical
    issues, and optimize performance to uphold system integrity and reliability. He is also
    experienced in providing comprehensive training and support to staff and end-users, empowering
    them to leverage Kronos systems effectively. He is an effective communicator with a
    collaborative approach, adept at working across departments to gather requirements, design
    solutions, and implement system enhancements. He possesses a proven ability to create detailed
    documentation, including reports, training materials, and system documentation, to facilitate
    smooth operations and knowledge transfer.
    Mr. Rahman's keen attention to detail, comprehensive skill set, commitment to continuous
    improvement, and proven track record make him a valuable asset in driving efficiency and
    compliance within Montgomery County's workforce management processes.
    """
    prompt = f""" 
    ""{resume}""
    ""{job_description}""
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": instruction},
            {"role": "user", "content": prompt}
        ]
    )
    total_tokens = response.usage.total_tokens
    print(total_tokens)
    return response.choices[0].message.content

# SOS generation: 
#   Paragraph 1: "Technuf provides # highly qualified candidates with strong experience in [job]"
#   Paragraph 2: Candidate capabilities and adeptness
#   Paragraph 3: Technuf (?) team ablities and how they help us perform the job well
def generate_sos(job_description, data, num, job_title):
    # Generates a scope of services given a job_description string, a data dict, a num int, and a job title string
    instruction = f""" 
    You are writing an understanding of the scope of services given a job description delimited by double quotes based on the abilities of the candidates given their summaries delimited by double quotes. 
    You will only reply in the following format totaling 3 paragraphs, do not include any unecessary greetings or introductions:
    The first paragraph will talk about the candidates' expertise on {job_title} should start with this sentence: Technuf provides {num} highly qualified candidates with strong experience in {job_title}.
    The second paragraph will talk about the candidates' capabilities and adeptness at {job_title}.
    The third paragraph will talk about the team's ability to fulfill the requirements of the role effectively. 
    """
    prompt = f""" 
    ""{job_description}""
    ""{data}""
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": instruction},
            {"role": "user", "content": prompt}
        ]
    )
    total_tokens = response.usage.total_tokens
    print(total_tokens)
    return response.choices[0].message.content

def get_title_page_info(text):
    # Extracts all necessary information from the text string and stores it in a dict
    title_info = {}
    topr_num_index = text.find("TOPR Number:")
    dept_index = text.find("Using Department:")
    tom_arr = text[text.find("Tech Contact:") + 13:text.find("TOPR Release")].strip().split("\n")
    tom_arr = list(filter(None, tom_arr))
    while ' ' in tom_arr:
        tom_arr.remove(' ')
        
    topr_title_arr = text[text.find("Order Title:") + 13 : dept_index].split("\n")
    topr_title_arr = list(filter(None, topr_title_arr))
    while ' ' in topr_title_arr:
        topr_title_arr.remove(' ')
        
    title_info['topr_num'] = text[topr_num_index + 13 : dept_index].strip().split()[0]
    title_info['dept'] = text[dept_index + 18 : text.find("Reference Contract:")].strip()
    title_info['due_date'] = text[text.find("Closing Date/Time:") + 18 : text.find("Approximate")].strip()
    title_info['submit_to_name'] = tom_arr[0].strip()
    title_info['submit_to_email'] = tom_arr[1].strip() if '@' in tom_arr[1] else tom_arr[2].strip()
    title_info['cat_type'] = "LCATS-3" if "LCATS" in text[:30] else "MCCATS-3"
    title_info['topr_title'] = "LSBRP Consulting and Technical Services 3\n(LCATS-3)" if "LCATS" in text[:30] else "Montgomery County Consulting and Technical Services 3\n(MCCATS-3)"
    title_info['job_title'] = topr_title_arr[0].strip()
    if len(topr_title_arr) > 1:
        for n in range(1, len(topr_title_arr)):
            title_info['job_title'] += " -- " + topr_title_arr[n].strip()
    return title_info

def format_resume(resume):
    # Takes a resume string and reformats that resume to the format that Technuf uses
    instruction = f""" 
    You are a resume formatter at Technuf LLC that needs to take a given resume delimited by double quotes and return it in the proper formatting shown below in brackets. Do not include the brackets in the actual response. Each section will be separated with two newlines so that there are 5 sections total: name, qualifications, education, skills, work experience, do not change this. You will only reply in the following format, do not include any unecessary greetings or introductions: 
    [Name of person on resume]
    \n\n
    [Give a summary of all of the resume's qualifications, knowledge, and experience in at least 15 bullet points:  • ]
    \n\n
    [Copy all of the resume's education, make sure to include all education information. If there is more than one education, include it on a new line with a bullet point]
    \n\n
    [Give a comma delimited list of all of the resume's skills]
    \n\n
    [Give a comprehensive and complete list of all of the resume's work experiences in this formatted example. Use this format to write out all of the resume's work experiences: Fannie Mae \t Jul 2016 - Dec 2019 \n Senior Agile Consultant / Project Management \n • Supported 3 Fannie Mae Single Family Portfolio teams with a corporate functions portfolio, 7 Enterprise Risk Management teams with Enterprise Finance Modeling, 1 CCFA team, 1 Governance, Risk, and Compliance team, and 3 Loan Quality Connect System / Loan Review Integration System teams at various Agile Maturity levels. \n • Supported 3 Fannie Mae Single Family Portfolio teams with a corporate functions portfolio, 7 Enterprise Risk Management teams with Enterprise Finance Modeling, 1 CCFA team, 1 Governance, Risk, and Compliance team, and 3 Loan Quality Connect System / Loan Review Integration System teams at various Agile Maturity levels.]
    """
    prompt = f"""
    ""{resume}""
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": instruction},
            {"role": "user", "content": prompt}
        ]
    )
    total_tokens = response.usage.total_tokens
    print(total_tokens)
    return response.choices[0].message.content

def generate_interview_questions(job_description, resume):
    # Generates interview questions and answers given a job_description string and a resume string
    skills_index = job_description.index("Responsibilities:")
    end_index = job_description.index("Education")
    job_description = job_description[skills_index:end_index - 1]

    prompt = f""" 
    What is the name on this given resume? {resume}
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You will only answer the question asked with only the answer. Do not include any filler words or greetings"},
            {"role": "user", "content": prompt}
        ]
    )
    name = response.choices[0].message.content
    
    instruction = f"""
    You are an interviewer at Technuf LLC that needs to curate interview questions to ask a candidate based on their resume and a given job description.
    According to this job description:
    {job_description}. 
    """
    prompt = f""" 
    Generate 10 interview questions and answers about the requirements section of the job description and 5 questions and answers tailored for this resume in bullet point format: 
    {resume}
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": instruction},
            {"role": "user", "content": prompt}
        ]
    )
    total_tokens = response.usage.total_tokens
    print(total_tokens)
    return (response.choices[0].message.content, name)

# TODO: Also add LinkedIn job posting here
def generate_job_description(text):
    # Generates a job description string in the Technuf format given a text string
    instruction = f"""
    You are a recruiter working at Technuf LLC. Only answer with the template given. Ensure to exclude any mention of Evaluation Criteria, How to Apply, Note, Contractor, Contract Staff, and Interview Criteria.
    You are tasked with generating a job posting following the format given, using information based on the solicitation requirements given. Reply in the following format, making sure that each bullet point should be on a new line, and every section should be separated with a new line, and each colon should have a new line after it: 
    Who we are [Do not change this header]:
    Technuf, LLC is a Maryland based SBA certified 8(a) small business company providing leading-edge and proven technologies, industry vertical domain expertise and highly skilled and motivated professionals to achieve our customers' mission critical business needs.
   
    What we're looking for [Do not change this header]:\n
    [Insert a one paragraph description of the name of the job that Technuf is looking for and what the job entails (make sure to start the paragraph with: Technuf is looking for...)]
    \n\n
    Responsibilities [Do not change this header]:\n
    [This section should copy everything from the Scope of Work section, but no not include anything that has the words "Contractor". Separate each bullet point with a new line]
    \n\n
    Skills and Experience [Do not change this header]:\n
    [This section should copy everything from the Skills/Experience section. Separate bullet points with a new line]
    \n\n
    Good to Have [Do not change this header]:\n
    [This section should copy everything from the Additionally, to be successful section. If there is no Additionally, to be successful section, then do not include the Good to Have section. Separate bullet points with a new line.]
    \n\n
    Education [Do not change this header]:\n
    Required: [insert required level of education here based on the job]\n
    Preferred: [insert preferred level of education here based on the job]\n
    """
    
    prompt = f""" 
    Use the information from the solicitations below to fill in the relevant sections:
    {text}
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": instruction},
            {"role": "user", "content": prompt}
        ]
    )
    total_tokens = response.usage.total_tokens
    print(total_tokens)
    return response.choices[0].message.content

def fill_document(template_path, data):
    # Fills a word document template with the data dict given a template_path, and lets you download the docx file
    doc = DocxTemplate(template_path)
    # This code prevents XML errors since docxtpl can't read in '&' symbols
    for key in data:
        data[key] = data[key].replace("&", "and")
        
    doc.render(data)
    

    # Makes the docx file downloadable
    bio = io.BytesIO()
    doc.save(bio)
    
    if doc:
        st.download_button(
            label="Click here to download",
            data=bio.getvalue(),
            file_name="proposal.docx",
            mime="docx"
        )
    
    # My attempt at allowing users to view the pdf file without having to download it. The problem with this is that 
    # it doesn't keep the pdf format
    bio.seek(0)
    result = mammoth.convert_to_html(bio)
    pdf = pdfkit.from_string(result.value)
    pdf_viewer(pdf, height=500)
    
def download_interview_questions(text, name):
    # Option to download interview questions generated given a text string
    doc = Document()
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    
    if doc:
        st.download_button(
            label="Click here to download",
            data=bio.getvalue(),
            file_name= name + "'s Interview Q and A.docx",
            mime="docx"
        )
def generate_proposal(resumes, concise_job_description, job_title):
    # Generates the proposal given up to 3 resume strings, a concise_job_description string, and a job_title string
    # This returns a data dict with all of the necessary information 
    concise_data = ""
    resume_count = 0
    for resume in resumes:
        formatted_resume = format_resume(resume)
        resume_list = formatted_resume.split("\n\n\n")
        resume_list = list(filter(None, resume_list))
        while ' ' in resume_list:
            resume_list.remove(' ')
        if resume_count == 0:
            data = {
                'resume_name_1' : resume_list[0].lstrip().rstrip(), 
                'resume_qualifications_1' : resume_list[1].lstrip().rstrip(), 
                'resume_education_1' : resume_list[2].lstrip().rstrip(), 
                'resume_skills_1' : resume_list[3].lstrip().rstrip()
            }
            work_str = ""
            for work in resume_list[4:]:
                work_str += work.lstrip().rstrip() + "\n"
            data['resume_work_1'] = work_str
                
            concise_data += data["resume_name_1"] + ": " + data["resume_qualifications_1"] + "\n" + data["resume_skills_1"]
            data.update({'resume_description_1' : generate_resume_descriptions(resume, concise_job_description)})
        elif resume_count == 1:
            data.update({
                'resume_name_2' : resume_list[0].lstrip().rstrip(), 
                'resume_qualifications_2' : resume_list[1].lstrip().rstrip(), 
                'resume_education_2' : resume_list[2].lstrip().rstrip(), 
                'resume_skills_2' : resume_list[3].lstrip().rstrip()
            })
            work_str = ""
            for work in resume_list[4:]:
                work_str += work.lstrip().rstrip()
            data['resume_work_2'] = work_str
            concise_data += data["resume_name_2"] + ": " + data["resume_qualifications_2"] + "\n" + data["resume_skills_2"]
            data.update({'resume_description_2' : generate_resume_descriptions(resume, concise_job_description)})
        elif resume_count == 2:
            data.update({
                'resume_name_3' : resume_list[0].lstrip().rstrip(), 
                'resume_qualifications_3' : resume_list[1].lstrip().rstrip(), 
                'resume_education_3' : resume_list[2].lstrip().rstrip(), 
                'resume_skills_3' : resume_list[3].lstrip().rstrip()
            })
            work_str = ""
            for work in resume_list[4:]:
                work_str += work.lstrip().rstrip()
            data['resume_work_2'] = work_str
            concise_data += data["resume_name_3"] + ": " + data["resume_qualifications_3"] + "\n" + data["resume_skills_3"]
            data.update({'resume_description_3' : generate_resume_descriptions(resume, concise_job_description)})
        resume_count += 1
                        
    data.update({
        'sos_understanding' : generate_sos(concise_job_description, concise_data, resume_count, job_title)
    })
    return data

# TODO: Button for unformatted proposal response, formatted proposal response, input resume to format it
# Focus on input resume + linkedin for now
def main():
    # Main function that has a login page and all of the streamlit elements
    hashed_passwords = Hasher(['Technuf123@']).generate()
    credentials = {
    "credentials": {
      'usernames': {
        "proposals": {
          "email": "proposals@technuf.com",
          "failed_login_attempts": 0,
          "logged_in": False,
          "name": "BD Proposal Team",
          "password": hashed_passwords[0]
        },
        "bdproposals": {
          "email": "bdproposals@technuf.com",
          "failed_login_attempts": 0,
          "logged_in": False,
          "name": "BD Proposal Team",
          "password": hashed_passwords[0]
        }
      }
    },
    "cookie": {
      "expiry_days": 30,
      "key": "some_signature_key",
      "name": "some_cookie_name"
    },
    "pre-authorized": {
      "emails": [
        "proposals@technuf.com"
      ]
    }
  }
    
    authenticator = stauth.Authenticate(credentials['credentials'],'some_cookie_name','some_signature_key',cookie_expiry_days=30,pre_authorized=["proposals@technuf.com"])
    # authenticator = stauth.Authenticate(
    #     config["credentials"],
    #     config["cookie"]["name"],
    #     config["cookie"]["key"],
    #     config["cookie"]["expiry_days"],
    #     config["pre-authorized"]
    # )
    authenticator.login()
    
    if st.session_state['authentication_status']:
        col1, col2, col3 = st.columns([26,26,8])
        with col3:
            authenticator.logout()
        data = {}
        # Add Technuf Logo and Technufbot text
        title_col1, mid, title_col2 = st.columns([1,2,20])
        with title_col1:
            st.image('technuf_logo.png', width=100)
        with title_col2:
            st.title("Technuf Proposal Bot")
        st.write("This proposal bot automates the initial process of breaking down the solicitation file that the proposal team receives, speeding up the entire process. The bot automatically extracts the scope of work, and other sections that are necessary to create a job description. The first button takes that scope of work and creates a job description based on that. The second button generates up to 10 interview questions for any one candidate. The third button generates an entire proposal response given a response template and the resumes of the candidates.")
        # Button session states
        if "last_pressed_button" not in st.session_state:
            st.session_state.last_pressed_button = None
        if 'clicked' not in st.session_state:
            st.session_state.clicked = False
        if 'disabled' not in st.session_state:
            st.session_state.disabled = True
        if st.session_state.get("enable_buttons", False):
            st.session_state.disabled = False
        job_description = ""
        job_title = ""
        pdf_text = ""
        concise_job_description = ""
        st.subheader("Choose a solicitation PDF or docx file")
        solicitation_file = st.file_uploader("a", type=[".pdf", ".docx"], label_visibility="collapsed")
        if solicitation_file:
            button_col1, button_col2, button_col3 = st.columns([18,20,20])
            with button_col1:
                job_desc_button = st.button("Generate Job Description", key='enable_buttons')
            with button_col2:
                interview_question_button = st.button("Generate Interview Questions")
            with button_col3:
                resume_upload_btn = st.button("Generate Proposal")
            # Parses through based on .pdf or .docx file
            if solicitation_file.name.endswith('.pdf'):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                    temp_file.write(solicitation_file.read())
                    temp_filepath = temp_file.name
                pdf_text = extract_text_from_pdf(temp_filepath)
                job_title = get_job_title_pdf(pdf_text)
                st.session_state.job_title = job_title
            elif solicitation_file.name.endswith('.docx'):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                    temp_file.write(solicitation_file.getbuffer())
                    temp_filepath = temp_file.name
                pdf_text = extract_text_from_docx(temp_filepath)
                job_title = get_job_title_docx(pdf_text)
                st.session_state.job_title = job_title
            # If scope of work is found, it will extract it and print it while filtering out the words contract staff
            relevant_sections = extract_relevant_sections(pdf_text)
            if relevant_sections:
                st.session_state.last_pressed_button = "job_desc"
                job_description = ""
                relevant_lines = relevant_sections.split('\n')
                filtered_lines = [line for line in relevant_lines if not ('contract staff' in line.lower())]
                relevant_sections = '\n'.join(filtered_lines)
                st.session_state['display_text'] = relevant_sections
                st.session_state['text_area_label'] = "Extracted Relevant Sections:"
                if job_desc_button:
                    st.session_state.job_button = True
                    with st.spinner('Generating job posting'):
                        job_posting = ""
                        if "job_posting" not in st.session_state:
                            job_posting = st.session_state.job_title + "\n\n"
                            job_description = generate_job_description(relevant_sections)
                            # Checks to see if there is an extra Note added on the end, and will remove it if there is
                            if job_description.find("Note:") != -1:
                                job_description = job_description[:job_description.find("Note:") + 1]
                            # st.subheader("Generated Job Description")
                            job_posting += job_description
                            job_posting += "\n\nBenefits:\n\nWe offer a competitive pay and benefits package that includes generous paid-time-off including holidays, short-and-long-term disability; group health insurance including medical, dental and vision coverage, training and 401(k) retirement plan. \n"
                            job_posting += "\nTechnuf is an Equal Opportunity/Affirmative Action Employer. Members of ethnic minorities, women, special disabled veterans, veterans of the Vietnam-era, recently separated veterans, and other protected veterans, persons of disability and/or persons age 40 and over are encouraged to apply.\n"
                            newline_index = job_posting.find("Preferred")
                            if newline_index:
                                job_posting = job_posting[:newline_index] + "\n" + job_posting[newline_index:]
                            st.session_state['job_posting'] = job_posting
                        else: 
                            job_posting = st.session_state.job_posting
                       
                        st.session_state['text_area_label'] = "Generated Job Description:"
                        st.session_state['display_text'] = job_posting
            else:
                st.error("No relevant sections found in the uploaded PDF.")
            # Upload resume to generate interview questions
            
            if interview_question_button:
                st.session_state.last_pressed_button = "interview"
            elif resume_upload_btn:
                st.session_state.last_pressed_button = "resume_upload"
            if st.session_state.last_pressed_button == "interview":
                if "job_posting" not in st.session_state:
                    with st.spinner('Currently loading...'):
                        job_posting = st.session_state.job_title + "\n\n"
                        job_posting += generate_job_description(pdf_text)
                        job_posting += "\n\nBenefits:\n\nWe offer a competitive pay and benefits package that includes generous paid-time-off including holidays, short-and-long-term disability; group health insurance including medical, dental and vision coverage, training and 401(k) retirement plan. \n"
                        job_posting += "\nTechnuf is an Equal Opportunity/Affirmative Action Employer. Members of ethnic minorities, women, special disabled veterans, veterans of the Vietnam-era, recently separated veterans, and other protected veterans, persons of disability and/or persons age 40 and over are encouraged to apply.\n"
                        newline_index = job_posting.find("Preferred")
                        if newline_index:
                            job_posting = job_posting[:newline_index] + "\n" + job_posting[newline_index:]
                        st.session_state.job_posting = job_posting
                st.subheader("Select the resume to generate 10 interview questions for:")
                resume_file = st.file_uploader("Select a resume to generate interview questions", type=[".pdf", ".docx"], accept_multiple_files=False, label_visibility="collapsed")
                if resume_file:
                    gen_int_button = st.button("Generate Interview Questions and Answers")
                    if gen_int_button:
                        with st.spinner('Generating Interview Questions and Answers...'):
                            if resume_file.name.endswith('.docx'):
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                                    temp_file.write(resume_file.getbuffer())
                                    resume_temp_file_path = temp_file.name
                                resume = extract_text_from_docx(resume_temp_file_path)
                            elif resume_file.name.endswith('.pdf'):
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                                    temp_file.write(resume_file.getbuffer())
                                    resume_temp_file_path = temp_file.name
                                resume = extract_text_from_pdf(resume_temp_file_path)
                            (interview_questions, name) = generate_interview_questions(st.session_state['job_posting'], resume)
                            st.session_state.display_text = interview_questions
                            st.session_state.text_area_label = "Interview Questions:"
                            download_interview_questions(interview_questions, name)
            # Format resume in the Technuf proposal format and add it into the template
            elif st.session_state.last_pressed_button == "resume_upload":
                if "job_posting" not in st.session_state:
                    with st.spinner('Currently loading...'):
                        job_posting = st.session_state.job_title + "\n\n"
                        job_posting += generate_job_description(pdf_text)
                        job_posting += "\n\nBenefits:\n\nWe offer a competitive pay and benefits package that includes generous paid-time-off including holidays, short-and-long-term disability; group health insurance including medical, dental and vision coverage, training and 401(k) retirement plan. \n"
                        job_posting += "\nTechnuf is an Equal Opportunity/Affirmative Action Employer. Members of ethnic minorities, women, special disabled veterans, veterans of the Vietnam-era, recently separated veterans, and other protected veterans, persons of disability and/or persons age 40 and over are encouraged to apply.\n"
                        newline_index = job_posting.find("Preferred")
                        if newline_index:
                            job_posting = job_posting[:newline_index] + "\n" + job_posting[newline_index:]
                        st.session_state.job_posting = job_posting
                st.subheader("Select the propsal template that you want to fill out:")
                template_file = st.file_uploader("Select the propsal template", type=[".docx"], accept_multiple_files=False, label_visibility="collapsed")
                st.subheader("Select the resumes you are selecting for the proposal:")
                resume_files = st.file_uploader("Select the resumes", type=[".pdf", ".docx"], accept_multiple_files=True, label_visibility="collapsed")
                resumes = []
                if resume_files and template_file:
                    gen_prop_button = st.button("Generate")
                    if gen_prop_button:
                        with st.spinner('This may take a couple of minutes, so grab a cup of coffee for now!'):
                            for resume in resume_files:
                                if resume.name.endswith('.docx'):
                                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                                        temp_file.write(resume.getbuffer())
                                        resume_temp_file_path = temp_file.name
                                    resumes.append(extract_text_from_docx(resume_temp_file_path))
                                elif resume.name.endswith('.pdf'):
                                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                                        temp_file.write(resume.getbuffer())
                                        resume_temp_file_path = temp_file.name
                                    resumes.append(extract_text_from_pdf(resume_temp_file_path))
                            
                            concise_job_description = job_description[job_description.find("Responsibilities"):job_description.find("Education")]
                            data = generate_proposal(resumes, concise_job_description, job_title)
                            data.update(get_title_page_info(pdf_text))
                            fill_document(template_file, data)
                            data_str = ""
                            for key, value in data.items():
                                data_str += key + ": " + value + "\n\n"
                            st.session_state.display_text = data_str
                            st.session_state.text_area_label = "Proposal Preview:"
            # Text box that will be updated to display everything
            if not st.session_state.last_pressed_button == "resume_upload":
                st.text_area(label=st.session_state.text_area_label, key="display_text", height=350)
        else:
            st.session_state.disabled = True
        
    elif st.session_state["authentication_status"] is False:
        st.error('Username/password is incorrect')
    elif st.session_state["authentication_status"] is None:
        st.warning('Please enter your username and password')
 
if __name__ == "__main__":
    main()
